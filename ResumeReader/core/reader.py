import docx, re, pythoncom, nltk, pandas as pd, string, random, ResumeReader.core.pdf2txt as pdf2txt
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from nltk.corpus import stopwords
#win32com.client
stop = stopwords.words('english') #For Name Extraction

def preprocess():
    print('something')
    
def findName(target):
    #Extracting Name
    names = []
    document = ' '.join([i for i in target.split() if i not in stop])
    sentences = nltk.sent_tokenize(document)
    sentences = [nltk.word_tokenize(sent) for sent in sentences]
    sentences = [nltk.pos_tag(sent) for sent in sentences]
    for tagged_sentence in sentences:
        for chunk in nltk.ne_chunk(tagged_sentence):
            if type(chunk) == nltk.tree.Tree:
                if chunk.label() == 'PERSON':
                    names.append(' '.join([c[0] for c in chunk]))
    name = names[0] if names else 'Not Found'
    return name
    
def readDetails(newDocUrl):
    target = 'Blank Document'
    linkedinPat = re.compile(r'((http(s)?://)?(www.)?(linkedin.com)(/|\\)(.+))')
    emailPat = re.compile(r'((mailto:)?[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)')
    name = linkedinUrl = email = phone = 'Not Found'
    fontSizes = fonts = set()
    tableCount = imageCount = pageCount = charCount = lineCount = 0
    error = ''
    
    if newDocUrl.endswith('.pdf'):
            target, pageCount = pdf2txt.pdf_to_text(newDocUrl)
            name = findName(target)
            #Extracting Email and LinkedinUrl
            _linkedinUrl = re.search(linkedinPat, target)
            _email = re.search(emailPat, target)
            if _linkedinUrl:
                linkedinUrl = _linkedinUrl.group(1) 
            if _email:
                email = _email.group(1)
    elif newDocUrl.endswith('.docx'):
        doc = docx.Document(newDocUrl)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        target = ' '.join(text)
        name = findName(target) 
        target = '|'.join(text)
            
        #Extracting Email and LinkedinUrl
        rels = doc.part.rels
        for rel in rels:
            if rels[rel].reltype == RT.HYPERLINK:
                temp = rels[rel]._target
                if linkedinPat.match(temp):
                    linkedinUrl = temp
                if emailPat.match(temp):
                    email = temp
                    
        #Extracting Font Data
        n = docx.styles.style._NumberingStyle
        for style in doc.styles:
            if not isinstance(style,n):
                if style.font.size != None:
                    fontSizes.add(style.font.size.pt)
                if style.font.name != None:
                    fonts.add(style.font.name)
                    
        #Extracting Table and Image Count
        for table in doc.tables:
            tableCount += 1
        for shape in doc.inline_shapes:
            if shape.type == 3:
                imageCount += 1   
                
        #Extracting other counts
        '''pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        word = word.Documents.Open(newDocUrl)
        word.Repaginate()
        pageCount = word.ComputeStatistics(2)
        charCount = word.ComputeStatistics(5)
        lineCount = word.ComputeStatistics(1)
        word.Close()'''
    else:
        error = "Incorrect Resume Format: Only .docx & .pdf files are accepted."
        
    #Extracting Phone
    target = re.findall(re.compile(r'(?:(?:\+?([1-9]|[0-9][0-9]|[0-9][0-9][0-9])\s*(?:[.-]\s*)?)?(?:\(\s*([2-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9])\s*\)|([0-9][1-9]|[0-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9]))\s*(?:[.-]\s*)?)?([2-9]1[02-9]|[2-9][02-9]1|[2-9][02-9]{2})\s*(?:[.-]\s*)?([0-9]{4})(?:\s*(?:#|x\.?|ext\.?|extension)\s*(\d+))?'), target)
    if target:
        phone = ''.join(target[0])
        if len(phone) > 10:
            phone = '+' + phone

    data = {
        'Item': ['Name', 'Phone', 'Email', 'LinkedIn', 'LineCount', 'CharCount', 'PageCount', 'Fonts', 'Font Sizes', 'tableCount', 'imageCount'],
        'Values': [name, phone, email, linkedinUrl, lineCount, charCount, pageCount, fonts,fontSizes, tableCount, imageCount]
    }
    
    return name, email, phone, linkedinUrl, lineCount, charCount, pageCount, fonts, fontSizes, tableCount, imageCount, data, error
    
def generateFiles(data, location):
    df = pd.DataFrame.from_dict(data,orient='index')
    name = ''.join(random.choices(string.ascii_lowercase + string.digits, k=30))
    csvLink = '/temp/' + name + '.csv'
    xlsxLink = '/temp/' + name + '.xlsx'
    df.to_csv(location + csvLink,index=False,header=False)
    read_file = pd.read_csv(location + csvLink)
    read_file.to_excel (location + xlsxLink, index = None, header=True)
    return csvLink, xlsxLink